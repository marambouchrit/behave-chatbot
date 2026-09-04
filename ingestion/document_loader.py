"""
document_loader.py
==================
Chargement et extraction du texte brut depuis les guides utilisateurs BeHave.

Formats supportés :
  - DOCX : fichiers Word  (python-docx)
  - PDF  : fichiers PDF   (PyPDF2)
"""

import re
import unicodedata
import logging
from pathlib import Path
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# ── Patterns compilés une seule fois au chargement du module ─────────────────

_RE_TABLE_SEPARATOR = re.compile(r"\|[\s\-|]+\|$")
_RE_DECORATION_LINE = re.compile(r"[-_=•·\s]+$")   
_RE_PAGE_NUMBER     = re.compile(r"[Pp]age\s+\d+\s*(sur|of|/)\s*\d+\.?$")
_RE_SIRYOS_HEADER   = re.compile(
    r"(DOCUMENT DE TRAVAIL"
    r"|GUIDE D['\u2019]UTILISATEUR"
    r"|R\u00e9f\s*:\s*DTR-[A-Z0-9\-/\s]+"  
    r")$",
    re.IGNORECASE,                           
)
_RE_BOLD_MARKDOWN   = re.compile(r"\*{1,2}([^*]+)\*{1,2}")
_RE_MULTI_SPACES    = re.compile(r"[ \t]{2,}")


_MODULE_PATTERNS: list[tuple[str, str]] = [
    ("predictive",  "BeHave Predictive"),
    ("master",      "BeHave Master Data"),
    ("access",      "BeHave Access"),
    ("acces",       "BeHave Access"),
    ("extraction",  "BeHave SAP Extraction"),
    ("sap",         "BeHave SAP Extraction"),
    ("pbi",         "BeHave Analytics (Power BI)"),
    ("power",       "BeHave Analytics (Power BI)"),
]


@dataclass
class Document:
    """Unité de base du pipeline RAG : texte + métadonnées."""
    page_content: str
    metadata: dict = field(default_factory=dict)

    def __repr__(self) -> str:
        preview = self.page_content[:80].replace("\n", " ")
        source  = self.metadata.get("source", "?")
        return f"Document(source={source!r}, preview={preview!r}...)"


def clean_text(text: str) -> str:
    """
    Nettoie le texte extrait d'un PDF ou DOCX en une seule passe sur les lignes.
    Aucun re.sub global après le join — évite une deuxième traversée du texte.
    """
    if not text:
        return ""

    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    cleaned_lines: list[str] = []

    for line in text.split("\n"):
        line = line.strip()

        if not line:
            continue
        if _RE_TABLE_SEPARATOR.match(line):
            continue
        if _RE_DECORATION_LINE.match(line) and len(line) > 3:
            continue
        if _RE_PAGE_NUMBER.match(line):
            continue
        if _RE_SIRYOS_HEADER.match(line):
            continue

        line = _RE_BOLD_MARKDOWN.sub(r"\1", line)
        line = _RE_MULTI_SPACES.sub(" ", line)
        cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()



def _infer_module(filename: str) -> str:
    """Déduit le module BeHave depuis le nom de fichier via table de dispatch."""
    filename_lower = filename.lower()
    for keyword, module_name in _MODULE_PATTERNS:
        if keyword in filename_lower:
            return module_name
    return "BeHave (module inconnu)"



def load_docx(file_path: str) -> Document:
    """Charge un fichier .docx et retourne un Document nettoyé."""
    import docx as python_docx
    path = Path(file_path)
    doc  = python_docx.Document(str(path))

    text_parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                text_parts.append(" | ".join(row_cells))

    cleaned = clean_text("\n".join(text_parts))

    if not cleaned:
        logger.warning("Fichier DOCX vide ou sans texte extractible : %s", path.name)

    return Document(
        page_content=cleaned,
        metadata={
            "source":    path.name,
            "file_path": str(path),
            "file_type": "docx",
            "module":    _infer_module(path.name),
        },
    )


def load_pdf(file_path: str) -> Document:
    """Charge un fichier .pdf et retourne un Document nettoyé."""
    import PyPDF2
    path       = Path(file_path)
    text_parts: list[str] = []
    num_pages  = 0

    with open(str(path), "rb") as f:
        reader    = PyPDF2.PdfReader(f)
        num_pages = len(reader.pages)

        for page_num in range(num_pages):
            try:
                page_text = reader.pages[page_num].extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            except Exception as exc:
                logger.warning("Page %d ignorée (%s) : %s", page_num + 1, path.name, exc)

    cleaned = clean_text("\n\n".join(text_parts))

    if not cleaned:
        logger.warning("Fichier PDF vide ou sans texte extractible : %s", path.name)

    return Document(
        page_content=cleaned,
        metadata={
            "source":    path.name,
            "file_path": str(path),
            "file_type": "pdf",
            "num_pages": num_pages,
            "module":    _infer_module(path.name),
        },
    )

def load_txt(file_path: str) -> Document:
    """Charge un fichier .txt et retourne un Document nettoyé."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8", errors="ignore")
    cleaned = clean_text(text)

    if not cleaned:
        logger.warning("Fichier TXT vide ou sans texte extractible : %s", path.name)

    return Document(
        page_content=cleaned,
        metadata={
            "source":    path.name,
            "file_path": str(path),
            "file_type": "txt",
            "module":    _infer_module(path.name),
        },
    )


_LOADERS: dict = {
    ".docx": load_docx,
    ".pdf":  load_pdf,
    ".txt":  load_txt,
}


def load_document(file_path: str) -> Document:
    """Charge un document en choisissant le loader selon l'extension."""
    path   = Path(file_path)
    suffix = path.suffix.lower()
    loader = _LOADERS.get(suffix)

    if loader is None:
        supported = ", ".join(_LOADERS.keys())
        raise ValueError(
            f"Format non supporté : '{suffix}'. Formats acceptés : {supported}"
        )

    return loader(file_path)


def load_all_documents(documents_dir: str) -> list[Document]:
    """
    Charge tous les documents DOCX/PDF d'un répertoire.
    Tout passe par logger — aucun print() direct.
    """
    documents_path = Path(documents_dir)

    if not documents_path.exists():
        raise FileNotFoundError(f"Répertoire introuvable : {documents_dir}")
    if not documents_path.is_dir():
        raise NotADirectoryError(f"Ce chemin n'est pas un répertoire : {documents_dir}")

    files = sorted(
        f for f in documents_path.iterdir()
        if f.suffix.lower() in _LOADERS
    )

    if not files:
        logger.warning("Aucun fichier DOCX/PDF trouvé dans : %s", documents_dir)
        return []

    documents: list[Document] = []

    for file_path in files:
        logger.info("Chargement : %s", file_path.name)

        try:
            doc = load_document(str(file_path))
        except Exception as exc:
            logger.error("Erreur sur %s : %s", file_path.name, exc)
            continue

        char_count = len(doc.page_content)

        if char_count == 0:
            logger.warning("Ignoré (vide après nettoyage) : %s", file_path.name)
            continue

        logger.info(
            "  ✓ %s — %d caractères | %s",
            file_path.name, char_count, doc.metadata["module"],
        )
        documents.append(doc)

    return documents



if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s  %(message)s")

    docs_dir = Path(__file__).parent.parent / "data" / "documents"
    logger.info("Répertoire : %s", docs_dir)

    docs = load_all_documents(str(docs_dir))

    logger.info("Résumé : %d document(s) chargé(s)", len(docs))
    for doc in docs:
        logger.info(
            "  • %s | %d chars | %s",
            doc.metadata["source"], len(doc.page_content), doc.metadata["module"],
        )