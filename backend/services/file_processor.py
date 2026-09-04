"""
file_processor.py
==================
Extraction et chunking des fichiers joints (PDF, DOCX) pour indexation
dans le mini-RAG temporaire (collection ChromaDB isolée par chat_id).

Le support images a été retiré : seuls les documents texte sont acceptés.
Chaque fichier est extrait puis découpé en chunks (Document) prêts à être
indexés via BeHaveVectorStore.add_temp_chunks().
"""

from io import BytesIO

import PyPDF2
from docx import Document as DocxDocument
from fastapi import UploadFile

from document_loader import Document

MAX_FILE_SIZE_BYTES   = 5 * 1024 * 1024
MAX_FILES_PER_REQUEST = 3

CHUNK_SIZE_CHARS   = 800
CHUNK_OVERLAP_CHARS = 100
MAX_CHUNKS_PER_FILE = 40

SUPPORTED_MIME_TYPES = frozenset({
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
})


class FileProcessingError(Exception):
    pass


# ─── Extraction ────────────────────────────────────────────────────────────

def _extract_pdf_text(content: bytes) -> str:
    reader = PyPDF2.PdfReader(BytesIO(content))
    pages = [page.extract_text() for page in reader.pages if page.extract_text()]
    text = '\n\n'.join(pages).strip()

    if not text:
        raise FileProcessingError(
            "Impossible d'extraire le texte de ce PDF. "
            "Le fichier est peut-être scanné ou protégé."
        )
    return text


def _extract_docx_text(content: bytes) -> str:
    doc = DocxDocument(BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = '\n\n'.join(paragraphs).strip()

    if not text:
        raise FileProcessingError(
            "Impossible d'extraire le texte de ce document Word."
        )
    return text


def _validate_file(filename: str, mime_type: str, size: int) -> None:
    if mime_type not in SUPPORTED_MIME_TYPES:
        raise FileProcessingError(
            f"Type de fichier non supporté : {mime_type}. "
            f"Types acceptés : PDF, DOCX."
        )
    if size > MAX_FILE_SIZE_BYTES:
        raise FileProcessingError(
            f"Fichier trop volumineux : {filename}. Taille maximale : 5 MB."
        )


# ─── Chunking ──────────────────────────────────────────────────────────────

def _chunk_text(text: str, filename: str) -> list[Document]:
    """Découpe le texte en chunks avec chevauchement, sur des frontières de mots."""
    words = text.split()
    if not words:
        return []

    chunks: list[Document] = []
    start = 0
    chunk_index = 0
    overlap_words = max(1, CHUNK_OVERLAP_CHARS // 6)

    while start < len(words) and chunk_index < MAX_CHUNKS_PER_FILE:
        current: list[str] = []
        current_len = 0
        end = start

        while end < len(words) and current_len < CHUNK_SIZE_CHARS:
            current.append(words[end])
            current_len += len(words[end]) + 1
            end += 1

        chunks.append(Document(
            page_content=" ".join(current),
            metadata={
                "filename": filename,
                "source": filename,
                "chunk_index": chunk_index,
            },
        ))
        chunk_index += 1

        next_start = end - overlap_words
        start = next_start if next_start > start else end

    return chunks


# ─── Fonction publique ─────────────────────────────────────────────────────

async def process_uploaded_files(files: list[UploadFile]) -> list[Document]:
    """
    Reçoit une liste d'UploadFile FastAPI (PDF/DOCX uniquement) et retourne
    une liste de Document (chunks) prêts à être indexés via
    BeHaveVectorStore.add_temp_chunks().

    Lève FileProcessingError si un fichier est invalide, illisible, ou d'un
    type non supporté.
    """
    if not files:
        return []

    if len(files) > MAX_FILES_PER_REQUEST:
        raise FileProcessingError(
            f"Maximum {MAX_FILES_PER_REQUEST} fichiers par message."
        )

    all_chunks: list[Document] = []

    for upload in files:
        content = await upload.read()
        mime_type = upload.content_type or ''
        filename = upload.filename or 'fichier'

        _validate_file(filename, mime_type, len(content))

        text = (
            _extract_pdf_text(content)
            if mime_type == 'application/pdf'
            else _extract_docx_text(content)
        )
        all_chunks.extend(_chunk_text(text, filename))

    return all_chunks